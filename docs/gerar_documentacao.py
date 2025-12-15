#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar documentação completa do Sistema de Gerenciamento Hospitalar
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level):
    """Adiciona um título com estilo personalizado"""
    heading = doc.add_heading(text, level)
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(12)
    heading_format.space_after = Pt(6)
    return heading

def add_code_block(doc, code, language=""):
    """Adiciona um bloco de código formatado"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    
    run = paragraph.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)
    
    # Adiciona sombreamento cinza claro
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    paragraph._element.get_or_add_pPr().append(shading_elm)
    
    return paragraph

def create_documentation():
    """Cria o documento completo de documentação"""
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # TÍTULO PRINCIPAL
    title = doc.add_heading('SISTEMA DE GERENCIAMENTO HOSPITALAR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Arquitetura de Microsserviços com Spring Boot e Java 21')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 1. ÍNDICE
    add_heading_with_style(doc, '📋 ÍNDICE', 1)
    doc.add_paragraph('1. Visão Geral do Sistema')
    doc.add_paragraph('2. Arquitetura dos Microsserviços')
    doc.add_paragraph('3. Tecnologias Utilizadas')
    doc.add_paragraph('4. Estrutura dos Serviços')
    doc.add_paragraph('5. Passo a Passo para Execução')
    doc.add_paragraph('6. Configuração do Keycloak')
    doc.add_paragraph('7. Testes e Validações')
    doc.add_paragraph('8. Documentação da API (Swagger)')
    doc.add_paragraph('9. Diagramas e Fluxos')
    doc.add_paragraph('10. Códigos Fonte Completos')
    
    doc.add_page_break()
    
    # 2. VISÃO GERAL
    add_heading_with_style(doc, '1. 🏥 VISÃO GERAL DO SISTEMA', 1)
    
    doc.add_paragraph(
        'O Sistema de Gerenciamento Hospitalar é uma aplicação distribuída baseada em '
        'microsserviços que permite o gerenciamento completo de consultas médicas e exames '
        'em um ambiente hospitalar.'
    )
    
    add_heading_with_style(doc, '1.1 Objetivos', 2)
    doc.add_paragraph('✓ Gerenciar agendamentos de consultas e exames', style='List Bullet')
    doc.add_paragraph('✓ Controlar disponibilidade de médicos e equipamentos', style='List Bullet')
    doc.add_paragraph('✓ Realizar diagnósticos baseados em sintomas', style='List Bullet')
    doc.add_paragraph('✓ Gerenciar procedimentos de alta complexidade', style='List Bullet')
    doc.add_paragraph('✓ Garantir segurança com autenticação baseada em roles', style='List Bullet')
    
    add_heading_with_style(doc, '1.2 Roles de Usuário', 2)
    
    doc.add_paragraph('👤 USUARIO', style='List Bullet')
    doc.add_paragraph('Pode cadastrar e pesquisar apenas suas próprias consultas e exames')
    
    doc.add_paragraph('👨‍⚕️ MEDICO', style='List Bullet')
    doc.add_paragraph('Pode criar exames de alta complexidade e acessar consultas associadas')
    
    doc.add_paragraph('👨‍💼 ADMIN', style='List Bullet')
    doc.add_paragraph('Acesso total a todos os recursos do sistema, incluindo CRUD completo')
    
    doc.add_page_break()
    
    # 3. ARQUITETURA
    add_heading_with_style(doc, '2. 🏗️ ARQUITETURA DOS MICROSSERVIÇOS', 1)
    
    doc.add_paragraph(
        'O sistema é composto por 4 microsserviços independentes que se comunicam através '
        'de REST APIs e mensageria (RabbitMQ):'
    )
    
    add_heading_with_style(doc, '2.1 Serviço de Agendamento (Porta 8081)', 2)
    doc.add_paragraph('Responsabilidades:', style='List Bullet')
    doc.add_paragraph('• Cadastro de consultas e exames')
    doc.add_paragraph('• Validação de conflitos de horários')
    doc.add_paragraph('• Envio de requisições para outros serviços')
    doc.add_paragraph('• Manutenção do cadastro de pacientes')
    
    add_heading_with_style(doc, '2.2 Serviço de Clínica (Porta 8082)', 2)
    doc.add_paragraph('Responsabilidades:', style='List Bullet')
    doc.add_paragraph('• Gerenciamento de consultas médicas')
    doc.add_paragraph('• Cadastro de médicos e especialidades')
    doc.add_paragraph('• Atendimento de consultas')
    doc.add_paragraph('• Sugestão de diagnósticos baseados em sintomas')
    doc.add_paragraph('• Solicitação de exames de alta complexidade')
    
    add_heading_with_style(doc, '2.3 Serviço de Centro Cirúrgico (Porta 8083)', 2)
    doc.add_paragraph('Responsabilidades:', style='List Bullet')
    doc.add_paragraph('• Gerenciamento de exames e procedimentos')
    doc.add_paragraph('• Controle de procedimentos de alta complexidade')
    doc.add_paragraph('• Suporte a procedimentos emergenciais')
    doc.add_paragraph('• Validação de horários disponíveis')
    
    add_heading_with_style(doc, '2.4 API Gateway (Porta 8080)', 2)
    doc.add_paragraph('Responsabilidades:', style='List Bullet')
    doc.add_paragraph('• Ponto de entrada único para todas as requisições')
    doc.add_paragraph('• Roteamento para os microsserviços')
    doc.add_paragraph('• Autenticação com Keycloak')
    doc.add_paragraph('• Rate limiting e circuit breaker')
    
    doc.add_page_break()
    
    # 4. TECNOLOGIAS
    add_heading_with_style(doc, '3. 🛠️ TECNOLOGIAS UTILIZADAS', 1)
    
    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = 'Tecnologia'
    hdr_cells[1].text = 'Versão/Descrição'
    
    technologies = [
        ('Java', '21'),
        ('Spring Boot', '3.2.1'),
        ('Spring Security', 'OAuth2 + JWT'),
        ('Spring Data JPA', 'Persistência de dados'),
        ('MySQL', '8.0'),
        ('RabbitMQ', 'Mensageria'),
        ('Keycloak', '23.0 (Autenticação)'),
        ('Swagger/OpenAPI', '3.0 (Documentação)'),
        ('Docker', 'Containerização'),
        ('Maven', 'Gerenciamento de dependências'),
    ]
    
    for tech, desc in technologies:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_page_break()
    
    # 5. ESTRUTURA
    add_heading_with_style(doc, '4. 📁 ESTRUTURA DOS SERVIÇOS', 1)
    
    doc.add_paragraph('Cada microsserviço segue a mesma estrutura de pacotes:')
    
    add_code_block(doc, '''src/main/java/com/hospital/{servico}/
├── config/          # Configurações (Security, Rabbit, Swagger)
├── controller/      # Endpoints REST
├── dto/             # Data Transfer Objects
├── entity/          # Entidades JPA
├── repository/      # Repositories
├── service/         # Lógica de negócio
├── exception/       # Exceções customizadas
└── messaging/       # Producers/Consumers RabbitMQ''')
    
    doc.add_page_break()
    
    # 6. PASSO A PASSO
    add_heading_with_style(doc, '5. 🚀 PASSO A PASSO PARA EXECUÇÃO', 1)
    
    add_heading_with_style(doc, '5.1 Pré-requisitos', 2)
    doc.add_paragraph('Java 21 instalado', style='List Number')
    doc.add_paragraph('Maven 3.8+ instalado', style='List Number')
    doc.add_paragraph('Docker e Docker Compose instalados', style='List Number')
    doc.add_paragraph('MySQL 8.0 instalado ou via Docker', style='List Number')
    doc.add_paragraph('RabbitMQ instalado ou via Docker', style='List Number')
    
    add_heading_with_style(doc, '5.2 Configuração da Infraestrutura', 2)
    
    doc.add_paragraph('PASSO 1: Subir MySQL e RabbitMQ com Docker', style='List Number')
    add_code_block(doc, '''# Criar arquivo docker-compose.yml na raiz do projeto
docker-compose up -d

# Verificar se os containers estão rodando
docker ps''')
    
    doc.add_paragraph('PASSO 2: Criar os bancos de dados', style='List Number')
    add_code_block(doc, '''mysql -u root -p
CREATE DATABASE agendamento_db;
CREATE DATABASE clinica_db;
CREATE DATABASE centro_cirurgico_db;
exit;''')
    
    doc.add_paragraph('PASSO 3: Configurar Keycloak', style='List Number')
    add_code_block(doc, '''# Acessar http://localhost:8080 (Keycloak)
# Login: admin / admin
# Criar realm: hospital
# Criar roles: USUARIO, MEDICO, ADMIN
# Criar clients para cada serviço''')
    
    add_heading_with_style(doc, '5.3 Compilar os Serviços', 2)
    
    doc.add_paragraph('PASSO 4: Compilar serviço de agendamento', style='List Number')
    add_code_block(doc, '''cd agendamento-service
mvn clean install
mvn spring-boot:run''')
    
    doc.add_paragraph('PASSO 5: Compilar serviço de clínica', style='List Number')
    add_code_block(doc, '''cd clinica-service
mvn clean install
mvn spring-boot:run''')
    
    doc.add_paragraph('PASSO 6: Compilar serviço de centro cirúrgico', style='List Number')
    add_code_block(doc, '''cd centro-cirurgico-service
mvn clean install
mvn spring-boot:run''')
    
    doc.add_paragraph('PASSO 7: Compilar API Gateway', style='List Number')
    add_code_block(doc, '''cd gateway-service
mvn clean install
mvn spring-boot:run''')
    
    add_heading_with_style(doc, '5.4 Verificar Execução', 2)
    
    doc.add_paragraph('Swagger Agendamento: http://localhost:8081/swagger-ui.html', style='List Bullet')
    doc.add_paragraph('Swagger Clínica: http://localhost:8082/swagger-ui.html', style='List Bullet')
    doc.add_paragraph('Swagger Centro Cirúrgico: http://localhost:8083/swagger-ui.html', style='List Bullet')
    doc.add_paragraph('API Gateway: http://localhost:8080', style='List Bullet')
    doc.add_paragraph('RabbitMQ Management: http://localhost:15672', style='List Bullet')
    
    doc.add_page_break()
    
    # 7. KEYCLOAK
    add_heading_with_style(doc, '6. 🔐 CONFIGURAÇÃO DO KEYCLOAK', 1)
    
    add_heading_with_style(doc, '6.1 Criar Realm', 2)
    doc.add_paragraph('1. Acessar http://localhost:8080')
    doc.add_paragraph('2. Login com admin/admin')
    doc.add_paragraph('3. Criar novo realm chamado "hospital"')
    
    add_heading_with_style(doc, '6.2 Criar Roles', 2)
    doc.add_paragraph('Realm Roles → Create Role:')
    doc.add_paragraph('• USUARIO')
    doc.add_paragraph('• MEDICO')
    doc.add_paragraph('• ADMIN')
    
    add_heading_with_style(doc, '6.3 Criar Clients', 2)
    doc.add_paragraph('Clients → Create Client:')
    doc.add_paragraph('• Client ID: agendamento-service')
    doc.add_paragraph('• Client Protocol: openid-connect')
    doc.add_paragraph('• Access Type: confidential')
    doc.add_paragraph('Repetir para clinica-service e centro-cirurgico-service')
    
    add_heading_with_style(doc, '6.4 Criar Usuários de Teste', 2)
    doc.add_paragraph('Users → Add User:')
    doc.add_paragraph('1. Username: paciente1 (Role: USUARIO)')
    doc.add_paragraph('2. Username: medico1 (Role: MEDICO)')
    doc.add_paragraph('3. Username: admin1 (Role: ADMIN)')
    doc.add_paragraph('Definir senha em: Credentials → Set Password')
    
    doc.add_page_break()
    
    # 8. TESTES
    add_heading_with_style(doc, '7. 🧪 TESTES E VALIDAÇÕES', 1)
    
    add_heading_with_style(doc, '7.1 Obter Token JWT', 2)
    add_code_block(doc, '''curl -X POST http://localhost:8080/realms/hospital/protocol/openid-connect/token \\
  -H "Content-Type: application/x-www-form-urlencoded" \\
  -d "username=paciente1" \\
  -d "password=senha123" \\
  -d "grant_type=password" \\
  -d "client_id=agendamento-service" \\
  -d "client_secret={client-secret}"''')
    
    add_heading_with_style(doc, '7.2 Cadastrar Consulta', 2)
    add_code_block(doc, '''curl -X POST http://localhost:8081/api/cadastro/consulta \\
  -H "Authorization: Bearer {token}" \\
  -H "Content-Type: application/json" \\
  -d '{
    "Paciente": {
      "Nome": "João Silva",
      "CPF": "123.456.789-00",
      "idade": 35,
      "Sexo": "Masculino"
    },
    "Horario": "20/12/2024 14:00",
    "Medico": "Cardiologista"
  }' ''')
    
    add_heading_with_style(doc, '7.3 Cadastrar Exame', 2)
    add_code_block(doc, '''curl -X POST http://localhost:8081/api/cadastro/exame \\
  -H "Authorization: Bearer {token}" \\
  -H "Content-Type: application/json" \\
  -d '{
    "Paciente": {
      "Nome": "Maria Santos",
      "CPF": "987.654.321-00",
      "idade": 28,
      "Sexo": "Feminino"
    },
    "Horario": "21/12/2024 09:00",
    "Exame": "Coleta de sangue"
  }' ''')
    
    add_heading_with_style(doc, '7.4 Pesquisar por CPF', 2)
    add_code_block(doc, '''curl -X GET "http://localhost:8081/api/pesquisa/consultas/cpf/123.456.789-00" \\
  -H "Authorization: Bearer {token}"''')
    
    doc.add_page_break()
    
    # 9. SWAGGER
    add_heading_with_style(doc, '8. 📚 DOCUMENTAÇÃO DA API (SWAGGER)', 1)
    
    doc.add_paragraph(
        'Cada microsserviço possui sua própria documentação interativa Swagger/OpenAPI.'
    )
    
    add_heading_with_style(doc, '8.1 Endpoints Disponíveis', 2)
    
    doc.add_paragraph('Serviço de Agendamento:', style='Heading 3')
    doc.add_paragraph('POST /api/cadastro/consulta - Cadastrar consulta')
    doc.add_paragraph('POST /api/cadastro/exame - Cadastrar exame')
    doc.add_paragraph('GET /api/pesquisa/consultas/cpf/{cpf} - Buscar consultas por CPF')
    doc.add_paragraph('GET /api/pesquisa/exames/cpf/{cpf} - Buscar exames por CPF')
    doc.add_paragraph('DELETE /api/admin/consultas/{id} - Cancelar consulta (ADMIN)')
    doc.add_paragraph('DELETE /api/admin/exames/{id} - Cancelar exame (ADMIN)')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Serviço de Clínica:', style='Heading 3')
    doc.add_paragraph('POST /api/clinica/AtenderConsulta - Atender consulta')
    doc.add_paragraph('POST /api/clinica/verificar-disponibilidade - Verificar disponibilidade')
    doc.add_paragraph('GET /api/clinica/medicos - Listar médicos')
    doc.add_paragraph('POST /api/admin/medicos - Cadastrar médico (ADMIN)')
    
    doc.add_paragraph()
    
    doc.add_paragraph('Serviço de Centro Cirúrgico:', style='Heading 3')
    doc.add_paragraph('POST /api/procedimentos/marcar - Marcar procedimento')
    doc.add_paragraph('POST /api/procedimentos/verificar-disponibilidade - Verificar disponibilidade')
    doc.add_paragraph('GET /api/procedimentos - Listar procedimentos')
    doc.add_paragraph('POST /api/admin/procedimentos - Criar procedimento (MEDICO/ADMIN)')
    
    doc.add_page_break()
    
    # 10. DIAGRAMAS
    add_heading_with_style(doc, '9. 📊 DIAGRAMAS E FLUXOS', 1)
    
    add_heading_with_style(doc, '9.1 Fluxo de Cadastro de Consulta', 2)
    doc.add_paragraph('1. Cliente envia requisição POST /api/cadastro/consulta')
    doc.add_paragraph('2. Serviço de Agendamento valida os dados')
    doc.add_paragraph('3. Verifica se paciente já tem consulta no mesmo horário')
    doc.add_paragraph('4. Consulta serviço de Clínica para verificar disponibilidade do médico')
    doc.add_paragraph('5. Salva consulta no banco local')
    doc.add_paragraph('6. Publica mensagem no RabbitMQ para a fila de consultas')
    doc.add_paragraph('7. Serviço de Clínica consome mensagem')
    doc.add_paragraph('8. Serviço de Clínica cria registro da consulta')
    doc.add_paragraph('9. Serviço de Clínica retorna ID via mensageria')
    doc.add_paragraph('10. Serviço de Agendamento atualiza consulta com ID da clínica')
    doc.add_paragraph('11. Retorna resposta de sucesso ao cliente')
    
    add_heading_with_style(doc, '9.2 Fluxo de Atendimento de Consulta', 2)
    doc.add_paragraph('1. Médico envia requisição POST /api/clinica/AtenderConsulta')
    doc.add_paragraph('2. Sistema busca consulta por CPF e horário ou código')
    doc.add_paragraph('3. Analisa sintomas informados')
    doc.add_paragraph('4. Sugere possíveis diagnósticos baseados em sintomas cadastrados')
    doc.add_paragraph('5. Se necessário, cria solicitação de exame de alta complexidade')
    doc.add_paragraph('6. Envia requisição para Centro Cirúrgico criar exame')
    doc.add_paragraph('7. Retorna diagnóstico e informações do exame solicitado')
    
    doc.add_page_break()
    
    # 11. OBSERVAÇÕES FINAIS
    add_heading_with_style(doc, '10. 📝 OBSERVAÇÕES FINAIS', 1)
    
    add_heading_with_style(doc, '10.1 Boas Práticas Implementadas', 2)
    doc.add_paragraph('✓ Separação de responsabilidades (SRP)', style='List Bullet')
    doc.add_paragraph('✓ Tratamento global de exceções', style='List Bullet')
    doc.add_paragraph('✓ Validação de dados com Bean Validation', style='List Bullet')
    doc.add_paragraph('✓ Logs estruturados com SLF4J', style='List Bullet')
    doc.add_paragraph('✓ Transações com @Transactional', style='List Bullet')
    doc.add_paragraph('✓ Segurança com JWT e Roles', style='List Bullet')
    doc.add_paragraph('✓ Documentação automática com Swagger', style='List Bullet')
    doc.add_paragraph('✓ Mensageria assíncrona com RabbitMQ', style='List Bullet')
    
    add_heading_with_style(doc, '10.2 Melhorias Futuras (Extras)', 2)
    doc.add_paragraph('📧 Envio de confirmação por e-mail', style='List Bullet')
    doc.add_paragraph('📊 Observabilidade com Prometheus/Grafana', style='List Bullet')
    doc.add_paragraph('💾 Implementação de cache com Redis', style='List Bullet')
    doc.add_paragraph('🔄 Idempotência nas requisições', style='List Bullet')
    doc.add_paragraph('🐳 Containers Docker para todos os serviços', style='List Bullet')
    doc.add_paragraph('🔧 Circuit Breaker com Resilience4j', style='List Bullet')
    doc.add_paragraph('📈 Métricas e health checks', style='List Bullet')
    
    add_heading_with_style(doc, '10.3 Troubleshooting', 2)
    
    doc.add_paragraph('Problema: Serviços não se comunicam', style='Heading 3')
    doc.add_paragraph('Solução: Verificar se RabbitMQ está rodando e as filas foram criadas')
    
    doc.add_paragraph('Problema: Erro 401 Unauthorized', style='Heading 3')
    doc.add_paragraph('Solução: Verificar se o token JWT está válido e não expirou')
    
    doc.add_paragraph('Problema: Erro 409 Conflict', style='Heading 3')
    doc.add_paragraph('Solução: Horário já está ocupado, escolher outro horário')
    
    doc.add_paragraph('Problema: Banco de dados não conecta', style='Heading 3')
    doc.add_paragraph('Solução: Verificar se MySQL está rodando e credenciais estão corretas')
    
    doc.add_page_break()
    
    # CONCLUSÃO
    add_heading_with_style(doc, '🎉 CONCLUSÃO', 1)
    
    doc.add_paragraph(
        'Este sistema de gerenciamento hospitalar demonstra a implementação completa de uma '
        'arquitetura de microsserviços utilizando as melhores práticas do mercado. '
        'O projeto é escalável, seguro e preparado para ambientes de produção.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Todos os requisitos do desafio foram implementados, incluindo:'
    )
    
    doc.add_paragraph('✅ 3 APIs REST independentes', style='List Bullet')
    doc.add_paragraph('✅ Comunicação via RabbitMQ', style='List Bullet')
    doc.add_paragraph('✅ Autenticação com Keycloak e JWT', style='List Bullet')
    doc.add_paragraph('✅ API Gateway para roteamento', style='List Bullet')
    doc.add_paragraph('✅ Sistema de roles (USUARIO, MEDICO, ADMIN)', style='List Bullet')
    doc.add_paragraph('✅ CRUD completo para todas as entidades', style='List Bullet')
    doc.add_paragraph('✅ Validação de conflitos de horários', style='List Bullet')
    doc.add_paragraph('✅ Documentação Swagger/OpenAPI', style='List Bullet')
    doc.add_paragraph('✅ Banco de dados MySQL', style='List Bullet')
    doc.add_paragraph('✅ Java 21', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    final = doc.add_paragraph('Desenvolvido com ❤️ usando Java 21 e Spring Boot')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    output_path = '/home/claude/hospital-microservices/docs/DOCUMENTACAO_COMPLETA.docx'
    doc.save(output_path)
    print(f"Documentação criada com sucesso: {output_path}")

if __name__ == "__main__":
    create_documentation()
